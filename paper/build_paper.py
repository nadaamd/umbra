"""
Génère le mini-papier de recherche Umbra (CBRI) au format Word (.docx).

- Équations rendues via matplotlib (mathtext) puis embarquées comme images nettes.
- Tableaux natifs Word.
- Les 3 figures réelles du backtest sont incluses.
Sortie : paper/Umbra_CBRI_Paper.docx
"""
import os
import struct
import tempfile

import matplotlib
matplotlib.use("Agg")
import matplotlib.pyplot as plt

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

HERE = os.path.dirname(os.path.abspath(__file__))
FIG_DIR = os.path.join(HERE, "..", "quant-backtest", "output")
OUT = os.path.join(HERE, "Umbra_CBRI_Paper.docx")
EQ_DIR = tempfile.mkdtemp(prefix="umbra_eq_")
EQ_DPI = 220
INK = RGBColor(0x14, 0x14, 0x14)
ACCENT = RGBColor(0x7A, 0x1F, 0x2B)   # bordeaux discret


# ── Helpers ──────────────────────────────────────────────────
def png_size(path):
    with open(path, "rb") as f:
        f.read(16)
        w, h = struct.unpack(">II", f.read(8))
    return w, h


def eq(latex, size=15):
    """Rend une équation LaTeX (mathtext) en PNG et renvoie son chemin."""
    # mathtext ne supporte pas certains raccourcis LaTeX -> conversion
    for a, b in [("\\big(", "\\left("), ("\\big)", "\\right)"),
                 ("\\Big(", "\\left("), ("\\Big)", "\\right)"),
                 ("\\big[", "\\left["), ("\\big]", "\\right]"),
                 ("\\Big[", "\\left["), ("\\Big]", "\\right]"),
                 ("\\!", ""), ("\\qquad", "\\ \\ "), ("\\quad", "\\ \\ "),
                 ("\\;", "\\,"), ("\\:", "\\,"), ("\\min", "\\mathrm{min}")]:
        latex = latex.replace(a, b)
    path = os.path.join(EQ_DIR, f"eq_{abs(hash(latex))}.png")
    fig = plt.figure(figsize=(0.01, 0.01))
    fig.text(0.5, 0.5, f"${latex}$", ha="center", va="center", fontsize=size, color="#141414")
    fig.savefig(path, dpi=EQ_DPI, bbox_inches="tight", pad_inches=0.06, transparent=True)
    plt.close(fig)
    return path


def add_eq(doc, latex, size=15):
    path = eq(latex, size)
    w, _ = png_size(path)
    width_in = min(6.3, w / EQ_DPI)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(8)
    p.add_run().add_picture(path, width=Inches(width_in))
    return p


def shade(p, fill="F3F1EC"):
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:fill"), fill)
    pPr.append(shd)


def callout(doc, label, text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.18)
    p.paragraph_format.right_indent = Inches(0.18)
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)
    r = p.add_run(f"{label} ")
    r.bold = True
    r.font.color.rgb = ACCENT
    r2 = p.add_run(text)
    r2.italic = True
    shade(p)
    return p


def body(doc, text):
    p = doc.add_paragraph(text)
    p.paragraph_format.space_after = Pt(6)
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    return p


def caption(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(text)
    r.italic = True
    r.font.size = Pt(9)
    r.font.color.rgb = RGBColor(0x55, 0x55, 0x55)
    p.paragraph_format.space_after = Pt(10)


def figure(doc, fname, cap, width=6.2):
    path = os.path.join(FIG_DIR, fname)
    if os.path.exists(path):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.add_run().add_picture(path, width=Inches(width))
        caption(doc, cap)


def table(doc, headers, rows):
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = "Light Grid Accent 1"
    for i, h in enumerate(headers):
        c = t.rows[0].cells[i]
        c.text = ""
        r = c.paragraphs[0].add_run(h)
        r.bold = True
        r.font.size = Pt(9.5)
    for row in rows:
        cells = t.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = ""
            r = cells[i].paragraphs[0].add_run(str(val))
            r.font.size = Pt(9.5)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)
    return t


# ── Document ─────────────────────────────────────────────────
doc = Document()

# marges
for s in doc.sections:
    s.top_margin = s.bottom_margin = Inches(0.9)
    s.left_margin = s.right_margin = Inches(1.0)

# styles
normal = doc.styles["Normal"]
normal.font.name = "Georgia"
normal.font.size = Pt(10.5)
normal.paragraph_format.line_spacing = 1.12
for hs in ("Title", "Heading 1", "Heading 2", "Heading 3"):
    st = doc.styles[hs]
    st.font.name = "Georgia"
    st.font.color.rgb = INK


def H1(t):
    h = doc.add_heading(t, level=1)
    h.paragraph_format.space_before = Pt(12)
    return h


def H2(t):
    h = doc.add_heading(t, level=2)
    h.paragraph_format.space_before = Pt(8)
    return h


# ── Titre ────────────────────────────────────────────────────
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = title.add_run("Umbra — A Composite Break-Risk Index (CBRI)")
r.bold = True
r.font.size = Pt(19)
r.font.name = "Georgia"
r.add_break()
r2 = title.add_run("for Autonomous DeFi Circuit-Breaking")
r2.bold = True
r2.font.size = Pt(19)
r2.font.name = "Georgia"

sub = doc.add_paragraph()
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
rs = sub.add_run("A quantitative early-warning score that detects systemic pool stress "
                 "and evacuates capital before liquidity evaporates.")
rs.italic = True
rs.font.size = Pt(11)

meta = doc.add_paragraph()
meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
rm = meta.add_run("Umbra Research  ·  ETH Global Lisbon 2026  ·  github.com/nadaamd/umbra")
rm.font.size = Pt(9.5)
rm.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

doc.add_paragraph()

# ── Abstract ─────────────────────────────────────────────────
ab = doc.add_paragraph()
rab = ab.add_run("Abstract. ")
rab.bold = True
ab.add_run(
    "DeFi liquidity pools fail fast: when a stablecoin de-pegs or liquidity flees, "
    "the price dislocates and exit costs explode within minutes, well before the "
    "average holder can react. We introduce the Composite Break-Risk Index (CBRI), "
    "a continuous 0–100 score that fuses three on-chain signals — the velocity of "
    "liquidity withdrawal, order-flow imbalance, and peg divergence — into a single "
    "risk reading. Each signal is squashed by a logistic function into a comparable "
    "[0,1] probability, then combined through a weighted Noisy-OR aggregator that "
    "behaves like a physical circuit breaker: the score trips if any single signal "
    "turns red, while non-informative signals are neutralised by a zero weight. We "
    "then derive the optimal trigger threshold τ* by maximising slippage-aware funds "
    "saved on real historical crashes. Back-tested on the March 2023 USDC de-peg "
    "(48,066 Uniswap v3 swaps, $5.57B volume), Umbra fires at τ*=66 seventeen hours "
    "before the trough — while USDC still trades at $1.0000 — preserving $126,900 "
    "(12.7%) on a $1M position.")
ab.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
ab.paragraph_format.left_indent = Inches(0.3)
ab.paragraph_format.right_indent = Inches(0.3)
ab.paragraph_format.space_before = Pt(4)
ab.paragraph_format.space_after = Pt(10)

# ── 1. Introduction ──────────────────────────────────────────
H1("1. The problem: DeFi breaks faster than people react")
body(doc,
     "A liquidity pool is a shared reserve of two assets that traders swap against. "
     "It works beautifully until confidence cracks. When holders rush to exit an "
     "asset — because a stablecoin looks unsafe, or a large lender is unwinding — two "
     "things happen at once: the price of the fleeing asset falls, and the depth of "
     "the pool (how much you can sell without moving the price) collapses. The result "
     "is a vicious spiral. By the time a retail user understands what is happening, "
     "the liquidity that would have let them exit cheaply is gone, and their sell "
     "order eats a punishing slippage on top of an already-fallen price.")
body(doc,
     "This is not hypothetical. The 2022–2023 cycle alone produced the Terra/UST "
     "collapse, the stETH discount, and the March 2023 USDC de-peg — each erasing "
     "hundreds of millions to tens of billions of dollars of value, much of it from "
     "users who simply reacted too late. The tooling gap is stark: markets have had "
     "circuit breakers for decades, yet DeFi positions have none.")
callout(doc, "In plain terms —",
        "Umbra is a smoke detector wired to a sprinkler. It watches the pool "
        "continuously, scores how close it is to breaking, and — past a proven "
        "threshold — automatically moves the user's funds to safety before the fire "
        "spreads.")

# ── 2. Design philosophy ─────────────────────────────────────
H1("2. What a good risk score must do")
body(doc,
     "Before any mathematics, we fixed three design principles that a credible "
     "break-risk score must satisfy:")
for lbl, txt in [
    ("Early, not coincident.",
     " The score must rise on the causes of a break (liquidity leaving, one-sided "
     "flow) — which appear first — not merely on the symptom (a price that has "
     "already fallen)."),
    ("Trip on any single failure.",
     " A pool can break through one channel alone. A score that averages its inputs "
     "dilutes a single extreme signal; a circuit breaker must not."),
    ("Calm must read calm.",
     " In a healthy market the score must sit near zero. A detector that cries wolf "
     "is worse than none — every false trigger costs the user real slippage."),
]:
    p = doc.add_paragraph(style="List Bullet")
    r = p.add_run(lbl)
    r.bold = True
    p.add_run(txt)

# ── 3. The three signals ─────────────────────────────────────
H1("3. The three signals")
body(doc,
     "Umbra reads three orthogonal on-chain quantities, each capturing a distinct "
     "way a pool tells you it is under stress. All are computed on 5-minute candles "
     "aggregated from raw swap, mint and burn events fetched via The Graph.")

H2("3.1  Liquidity-drain velocity — the early warning")
body(doc,
     "The single earliest tell of a crisis is not the price; it is liquidity "
     "providers pulling their capital out, and doing so ever faster. We measure the "
     "relative rate at which the pool's value is draining:")
add_eq(doc, r"v_L \;=\; -\,\frac{1}{L_t}\,\frac{\Delta L}{\Delta t}")
body(doc,
     "Here L is the pool's liquidity (total value locked) and ΔL/Δt its change per "
     "unit time. Dividing by L makes the signal scale-free: a 5%/hour drain means the "
     "same thing for a $1M pool and a $1B pool. We compute ΔL from the net of mint "
     "(liquidity added) and burn (liquidity removed) events over a 45-minute window, "
     "which smooths single-LP noise while remaining fast.")
callout(doc, "Why it leads —",
        "Sophisticated liquidity providers de-risk first. Their exit thins the book "
        "hours before the crowd notices, so a rising v_L is a genuine head start.")

H2("3.2  Order-flow imbalance — the rush to the exit")
body(doc,
     "In a panic, trading becomes one-directional: everyone sells the same asset. We "
     "quantify how lopsided the flow is:")
add_eq(doc, r"I \;=\; \frac{\left|\;\sum \mathrm{in} \;-\; \sum \mathrm{out}\;\right|}"
            r"{\sum \left|\,\mathrm{flow}\,\right|}")
body(doc,
     "I lives in [0,1]: it is 0 when buys and sells balance, and approaches 1 when "
     "flow is entirely one-way. Because Uniswap v3 has no static reserves, this "
     "realized order-flow imbalance is a more faithful stress measure for "
     "concentrated liquidity than any reserve-ratio proxy.")

H2("3.3  Peg divergence — the confirmation")
body(doc,
     "For a stablecoin, the most direct evidence of a break is the price itself "
     "leaving its peg. We read the price straight from the deepest stable/stable pool "
     "(USDC/USDT), which needs no external oracle:")
add_eq(doc, r"D \;=\; \left|\,1 - P_{\mathrm{USDC}}\,\right|")
body(doc,
     "D is the fraction by which the coin has left $1. It is a confirming signal — "
     "large and unambiguous once a de-peg is underway — that complements the "
     "forward-looking drain and flow signals.")

# ── 4. From signals to a score ───────────────────────────────
H1("4. From three signals to one 0–100 score")

H2("4.1  Normalization: making signals comparable")
body(doc,
     "The three raw quantities live on incompatible scales (a fraction per hour, a "
     "ratio, a price deviation). We map each onto a common [0,1] axis with a logistic "
     "(sigmoid) function, which acts as a soft, tunable threshold:")
add_eq(doc, r"s_i \;=\; \sigma\!\big(\alpha_i\,(x_i-\theta_i)\big),"
            r"\qquad \sigma(z)=\frac{1}{1+e^{-z}}")
body(doc,
     "θ_i is the level at which the signal starts to look abnormal, and α_i sets how "
     "sharply the transition happens. Below its threshold a signal contributes ≈0; "
     "well above it, ≈1. Each s_i can be read as “the probability, according to this "
     "one lens, that the pool is breaking.”")

H2("4.2  Aggregation: the circuit-breaker logic (Noisy-OR)")
body(doc,
     "We do not average the signals — averaging would let a single extreme reading be "
     "diluted by two calm ones, exactly the wrong behaviour for a breaker. Instead we "
     "combine them with a weighted Noisy-OR:")
add_eq(doc, r"\mathrm{CBRI} \;=\; 100\left(1-\prod_i\big(1-w_i\,s_i\big)\right)")
body(doc,
     "Read literally, the product is the probability that the pool survives every "
     "signal at once; one minus that is the probability it breaks through at least "
     "one channel. The consequence is precisely the circuit-breaker property: if any "
     "single s_i approaches 1, the whole score is driven to 100, regardless of the "
     "others. The weights w_i are a safety valve — a signal that carries no "
     "information on a given regime is set to w=0 and cleanly removed, without "
     "re-deriving the model.")
callout(doc, "A useful discovery —",
        "On the USDC de-peg, order-flow imbalance turned out to be non-discriminating "
        "(arbitrageurs kept buying the cheap coin, so flow never became fully "
        "one-way). Rather than let it add noise, we set its weight to zero. The CBRI "
        "here is effectively a two-signal breaker — drain OR de-peg:")
add_eq(doc, r"\mathrm{CBRI} \;=\; 100\left(1-(1-s_{\mathrm{drain}})"
            r"\,(1-s_{\mathrm{depeg}})\right)")

# ── 5. Calibration ───────────────────────────────────────────
H1("5. Calibration: calm reads ~8, crisis reads 100")
body(doc,
     "The thresholds θ_i and steepnesses α_i are not guessed; they are set so that "
     "each sigmoid rests near zero in normal conditions and saturates only in genuine "
     "stress, then validated against the data. With the parameters below, the CBRI "
     "sits at a resting baseline of roughly 7–10 during calm trading and rises to 100 "
     "at the height of the crisis — a clean separation with no false alarms in the "
     "quiet window.")
table(doc,
      ["Signal", "Threshold θ", "Steepness α", "Window", "Weight w"],
      [["Liquidity-drain velocity v_L", "0.06 /h", "60", "45 min", "1.0"],
       ["Order-flow imbalance I", "0.30", "15", "30 min", "0.0 (off)"],
       ["Peg divergence D", "0.012", "250", "—", "1.0"]])
caption(doc, "Table 1 — CBRI model parameters (5-minute candles).")

# ── 6. From score to action ──────────────────────────────────
H1("6. From score to action: the optimal trigger τ*")
body(doc,
     "A score is only useful if it tells you when to act. Umbra evacuates the "
     "position — swapping the at-risk asset into a safe stablecoin — the first time "
     "the CBRI crosses a threshold τ. Choosing τ well is itself a quantitative "
     "problem.")

H2("6.1  The false-recovery trap")
body(doc,
     "USDC eventually recovered to ~$1.00. Measured naïvely against the trough, "
     "evacuating always looks good and earlier is always better; measured against the "
     "full recovery, it looks pointless. Both framings are misleading. The honest "
     "objective treats evacuation as downside protection whose real cost is the "
     "slippage paid to get out — and that slippage is the crux.")

H2("6.2  Slippage-aware funds saved")
body(doc,
     "As liquidity collapses, the cost of exiting a fixed $1M position explodes. We "
     "model exit slippage as inversely related to the pool's surviving depth, "
     "anchored to a live Uniswap quote in calm conditions:")
add_eq(doc, r"\mathrm{slip}(t) \;=\; \min\!\Big(\mathrm{CAP},\;"
            r"s_{\mathrm{calm}}\,(L_{\mathrm{ref}}/L_t)^{\beta}\Big)")
body(doc,
     "The funds saved by triggering at threshold τ are then the value recovered at "
     "the exit — net of that slippage — versus suffering the trough:")
add_eq(doc, r"\mathrm{FundsSaved}(\tau) \;=\; P\cdot P_{\mathrm{exit}}(\tau)\,"
            r"\big(1-\mathrm{slip}(\tau)\big)\;-\;P\cdot P_{\mathrm{trough}}")
body(doc, "and the optimal trigger maximises this in expectation:")
add_eq(doc, r"\tau^{*} \;=\; \mathrm{arg\,max}_{\;\tau}\;"
            r"\mathrm{E}\big[\mathrm{FundsSaved}(\tau)\big]")

H2("6.3  The plateau and the chosen threshold")
body(doc,
     "Because the CBRI leaps from its resting baseline straight past 60 the moment "
     "liquidity flight begins, every threshold up to that jump triggers at the same "
     "optimal instant — a plateau of equally-good choices. We select the highest "
     "threshold on that plateau, τ*=66, which captures the full protection while "
     "leaving the maximum safety margin against false alarms.")

# ── 7. Empirical validation ──────────────────────────────────
H1("7. Empirical validation — the March 2023 USDC de-peg")
body(doc,
     "We replay the SVB-driven USDC de-peg of 10–13 March 2023 on real data: 48,066 "
     "Uniswap v3 swaps totalling $5.57B of volume, cross-referenced with the "
     "USDC/USDT pool for the peg and with mint/burn events for the liquidity drain. "
     "The evacuated position is $1M of USDC.")

figure(doc, "fig1_cbri_vs_price.png",
       "Figure 1 — The CBRI (red) against the USDC price (black). The breaker fires "
       "at τ*=66 on 10 March 14:10 UTC, while USDC is still $1.0000 — about 17 hours "
       "before the $0.8726 trough.")

body(doc,
     "The active liquidity of the USDC/USDT pool collapsed by a factor of roughly "
     "23 million as the price left the concentrated range — the physical reason exit "
     "slippage explodes if one waits. The table below shows the cost of hesitation: "
     "each block of thresholds triggers later, at a worse price and a worse slippage.")
table(doc,
      ["Trigger τ", "Exit time (UTC)", "Exit price", "Slippage", "Funds saved"],
      [["10–66  (τ*)", "10 Mar 14:10", "$1.0000", "5 bps", "$126,900"],
       ["67–72", "11 Mar 00:15", "$0.9892", "53 bps", "$111,400"],
       ["73–98", "11 Mar 01:00", "$0.9726", "91 bps", "$91,100"],
       ["99", "11 Mar 03:00", "$0.9371", "657 bps", "$2,900"]])
caption(doc, "Table 2 — Waiting is expensive: funds saved by trigger threshold, $1M position.")

figure(doc, "fig2_funds_saved_vs_tau.png",
       "Figure 2 — Funds saved as a function of the trigger threshold τ. The flat "
       "plateau (τ ≤ 66) is the zero-false-positive optimum; beyond it, every step of "
       "waiting forfeits capital.")

figure(doc, "fig3_slippage_explosion.png",
       "Figure 3 — Exit slippage over time versus the calm Uniswap anchor. Exiting at "
       "τ* costs a few basis points; waiting for price confirmation means hitting a "
       "liquidity wall.")

body(doc,
     "Headline result: triggering at τ*=66 evacuates the position at $1.0000 for a "
     "5-basis-point cost, preserving $126,900 — 12.7% of the position — that a "
     "passive holder would have watched evaporate into the trough.")

# ── 8. Economic alignment ────────────────────────────────────
H1("8. Economic alignment")
body(doc,
     "Umbra charges a success fee only on the loss it avoids. On the case above, a "
     "10% fee is $12,690 of revenue — earned strictly because the user kept "
     "$126,900 they would otherwise have lost. Incentives are exactly aligned: Umbra "
     "is paid if and only if the user is protected. Every score and the exact model "
     "that produced it are anchored on 0G decentralized storage, so the scoring is "
     "auditable and tamper-evident rather than a black box.")

# ── 9. Limitations ───────────────────────────────────────────
H1("9. Limitations and future work")
for lbl, txt in [
    ("Single-venue execution.",
     " The MVP routes the evacuation through one pool — the very pool whose depth "
     "collapses in a crisis. Production execution must split across pools, fee tiers "
     "and venues (multi-hop routing) to avoid the slippage wall it warns about."),
    ("Signal set.",
     " Order-flow imbalance was non-informative on this crash and disabled. A "
     "volatility-surge or cross-pool contagion signal is a natural third pillar for a "
     "broader regime library."),
    ("Recovery uncertainty.",
     " The τ* objective is honest downside protection; a fuller treatment weights "
     "terminal-collapse (UST-like) against false-alarm (USDC-like) scenarios under an "
     "explicit crisis probability."),
    ("Real-time pipeline.",
     " The current system replays history for validation; a production loop streams "
     "The Graph continuously and pushes scores on-chain in real time."),
]:
    p = doc.add_paragraph(style="List Bullet")
    r = p.add_run(lbl)
    r.bold = True
    p.add_run(txt)

# ── Appendix ─────────────────────────────────────────────────
H1("Appendix — Data and reproducibility")
body(doc,
     "Data source: The Graph (Uniswap v3 Ethereum mainnet subgraph). Pools: "
     "USDC/WETH 0.05% (0x88e6…5640) for drain and flow; USDC/USDT 0.01% (0x3416…27c6) "
     "for the peg. Window: 10–13 March 2023 UTC. Execution venue: Uniswap v3 "
     "(QuoterV2 + SwapRouter02). The full pipeline — extraction, CBRI computation, "
     "τ* back-test, and 29 unit/end-to-end tests — is reproducible from the "
     "repository via a single script (./e2e.sh).")

doc.add_paragraph()
foot = doc.add_paragraph()
foot.alignment = WD_ALIGN_PARAGRAPH.CENTER
rf = foot.add_run("Umbra · CBRI v1.0 · Composite Break-Risk Index · ETH Global Lisbon 2026")
rf.font.size = Pt(8.5)
rf.font.color.rgb = RGBColor(0x88, 0x88, 0x88)

doc.save(OUT)
print(f"✅ Papier généré : {OUT}")
print(f"   pages estimées : ~{len(doc.paragraphs)//18 + 1}  ·  équations & 3 figures embarquées")
